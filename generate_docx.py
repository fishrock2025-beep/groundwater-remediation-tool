#!/usr/bin/env python3
"""
地下水调查报告 Word文档生成工具
根据报告编制格式.docx规范，将txt文本转换为格式规范的Word文档

用法:
    python3 generate_docx.py 输入文件.txt -o 输出文件.docx
"""

import sys
import os
import re

sys.path.insert(0, '/tmp/docx_venv/lib/python3.12/site-packages')

try:
    from docx import Document
    from docx.shared import Pt, Cm, Emu
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
    from docx.oxml.ns import qn, nsdecls
    from docx.oxml import parse_xml
except ImportError:
    print("错误: 需要安装python-docx")
    print("请运行: python3 -m venv /tmp/docx_venv && /tmp/docx_venv/bin/pip install python-docx")
    sys.exit(1)


def set_font(run, font_cn='仿宋_GB2312', font_en='Times New Roman', size_pt=12, bold=False):
    """设置字体属性"""
    run.font.name = font_en
    run.font.size = Pt(size_pt)
    run.font.bold = bold
    r = run._element
    rPr = r.find(qn('w:rPr'))
    if rPr is None:
        rPr = parse_xml(f'<w:rPr {nsdecls("w")}></w:rPr>')
        r.insert(0, rPr)
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = parse_xml(f'<w:rFonts {nsdecls("w")}></w:rFonts>')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:eastAsia'), font_cn)


def set_pf(paragraph, line_spacing=1.5, indent=None, alignment=None):
    """设置段落格式"""
    pf = paragraph.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.line_spacing = line_spacing
    if indent is not None:
        pf.first_line_indent = indent
    if alignment is not None:
        paragraph.alignment = alignment


def setup_page(section):
    """设置页面（A4, 上3下3.5左3.17右3.17）"""
    section.top_margin = Cm(3)
    section.bottom_margin = Cm(3.5)
    section.left_margin = Cm(3.17)
    section.right_margin = Cm(3.17)
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)


def txt_to_docx(input_path, output_path=None):
    """将txt报告文件转换为格式规范的Word文档"""
    if output_path is None:
        output_path = os.path.splitext(input_path)[0] + '.docx'

    with open(input_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    # 跳过系统警告
    if lines and 'warning: setlocale' in lines[0]:
        lines = lines[1:]

    doc = Document()
    setup_page(doc.sections[0])

    # 设置默认样式
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(12)
    rPr = style.element.find(qn('w:rPr'))
    if rPr is not None:
        rFonts = parse_xml(f'<w:rFonts {nsdecls("w")} w:eastAsia="仿宋_GB2312"/>')
        rPr.append(rFonts)

    chapter_pat = re.compile(r'^第[一二三四五六七八九十百]+章\s+')
    section_pat = re.compile(r'^\d+\.\d+\s+')
    subsection_pat = re.compile(r'^\d+\.\d+\.\d+\s+')
    list_pat = re.compile(r'^（[一二三四五六七八九十百]+）|^（\d+）|^\(\d+\)')
    bullet_pat = re.compile(r'^- |^• |^  - ')

    for line in lines:
        line = line.strip()
        if not line:
            doc.add_paragraph()
            continue

        # 移除行号
        if '|' in line:
            parts = line.split('|', 1)
            if len(parts) > 1 and parts[0].strip().isdigit():
                line = parts[1].strip()

        if not line:
            continue

        if chapter_pat.match(line):
            h = doc.add_heading('', 1)
            run = h.add_run(line)
            set_font(run, '黑体', 'Times New Roman', 16, bold=True)
            set_pf(h, 1.5, alignment=WD_ALIGN_PARAGRAPH.LEFT)
            doc.add_page_break()

        elif section_pat.match(line):
            h = doc.add_heading('', 2)
            run = h.add_run(line)
            set_font(run, '仿宋_GB2312', 'Times New Roman', 15, bold=True)
            set_pf(h, 1.5, alignment=WD_ALIGN_PARAGRAPH.LEFT)

        elif subsection_pat.match(line):
            h = doc.add_heading('', 3)
            run = h.add_run(line)
            set_font(run, '楷体_GB2312', 'Times New Roman', 15, bold=True)
            set_pf(h, 1.5, alignment=WD_ALIGN_PARAGRAPH.LEFT)

        elif list_pat.match(line):
            p = doc.add_paragraph()
            run = p.add_run(line)
            set_font(run, '仿宋_GB2312', 'Times New Roman', 12)
            set_pf(p, 1.5, indent=Emu(355600))

        elif bullet_pat.match(line):
            p = doc.add_paragraph()
            run = p.add_run(re.sub(r'^- |^• |^  - ', '', line))
            set_font(run, '仿宋_GB2312', 'Times New Roman', 12)
            set_pf(p, 1.5, indent=Emu(355600))

        else:
            p = doc.add_paragraph()
            run = p.add_run(line)
            set_font(run, '仿宋_GB2312', 'Times New Roman', 12)
            set_pf(p, 1.5, indent=Emu(355600))

    doc.save(output_path)
    file_size = os.path.getsize(output_path)
    print(f"文档已生成: {output_path}")
    print(f"段落数: {len(doc.paragraphs)}, 文件大小: {file_size/1024:.1f} KB")
    return output_path


if __name__ == '__main__':
    import argparse
    parser = argparse.ArgumentParser(description='将地下水调查报告txt转换为格式规范的Word文档')
    parser.add_argument('input', help='输入的txt文件路径')
    parser.add_argument('-o', '--output', help='输出的docx文件路径（默认自动生成）')
    args = parser.parse_args()
    txt_to_docx(args.input, args.output)
